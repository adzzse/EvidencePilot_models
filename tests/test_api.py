import asyncio
import httpx
import io
import json
import logging
import zipfile
from types import SimpleNamespace
from uuid import uuid4
from zipfile import ZipFile

import pytest
from docx import Document as WordDocument
from fastapi.testclient import TestClient

import app.extraction as extraction
import app.main as main
from app import limits
from app.extraction import (
    ExtractedDocument,
    ExtractionError,
    ExtractionUnavailableError,
    ExtractionWorkProduct,
    extract_from_url,
)
from app.models import ExtractRequest, ExtractionBlock, GenerateResponse
from app.generation import (
    GenerationConfigurationError,
    GenerationInvalidResponseError,
    GenerationRateLimitError,
    GenerationUnavailableError,
    OllamaGenerationProvider,
    OpenAICompatibleGenerationProvider,
)
from app.settings import Settings


SETTINGS = Settings(
    generation_provider="ollama",
    model_api_key="test-key",
    extraction_allowed_hosts=("storage.test",),
)
HEADERS = {"X-API-Key": "test-key"}


@pytest.fixture(autouse=True)
def settings_override(monkeypatch):
    monkeypatch.setattr(limits, "local_gate", limits.ModelCallGate(4))
    monkeypatch.setattr(limits, "generation_gate", limits.ModelCallGate(4))
    main.app.dependency_overrides[main.get_settings] = lambda: SETTINGS
    try:
        yield
    finally:
        main.app.dependency_overrides.clear()


@pytest.fixture
def client() -> TestClient:
    return TestClient(main.app)


def test_model_call_gate_caps_concurrency_and_spaces_starts():
    async def exercise():
        gate = limits.ModelCallGate(max_concurrent=2, min_interval_ms=10)
        starts = []
        active = 0
        max_active = 0

        async def invoke():
            nonlocal active, max_active
            async with gate.slot():
                starts.append(asyncio.get_running_loop().time())
                active += 1
                max_active = max(max_active, active)
                await asyncio.sleep(0.1)
                active -= 1

        await asyncio.gather(*(invoke() for _ in range(3)))
        return starts, max_active

    starts, max_active = asyncio.run(exercise())

    assert max_active == 2
    assert all(later - earlier >= 0.007 for earlier, later in zip(starts, starts[1:]))


def test_health_degrades_without_taking_api_offline(client: TestClient, monkeypatch):
    async def unavailable(_):
        raise main.OllamaUnavailableError("offline")

    monkeypatch.setattr(OllamaGenerationProvider, "health", unavailable)
    response = client.get("/health")

    assert response.status_code == 200
    assert response.json()["status"] == "degraded"


def test_health_with_remote_generation_requires_only_local_embedding(
    client: TestClient, monkeypatch
):
    required_generation = []
    settings = SETTINGS.model_copy(update={
        "generation_provider": "remote",
        "generation_api_key": "secret",
        "generation_base_url": "https://gateway.test/v1",
        "generation_model": "test-model",
    })
    main.app.dependency_overrides[main.get_settings] = lambda: settings

    async def local_health(_, require_generation=True):
        required_generation.append(require_generation)
        return {
            "ok": True,
            "model_available": False,
            "embedding_model_available": True,
        }

    async def remote_health(_):
        return {
            "ok": True,
            "provider": "remote",
            "model": settings.generation_model,
        }

    monkeypatch.setattr(main, "check_ollama", local_health)
    monkeypatch.setattr(
        OpenAICompatibleGenerationProvider,
        "health",
        remote_health,
    )

    response = client.get("/health")

    assert response.status_code == 200
    assert response.json()["status"] == "ok"
    assert response.json()["generation_provider"] == "remote"
    assert response.json()["generation"] == {
        "ok": True,
        "provider": "remote",
        "model": settings.generation_model,
    }
    assert required_generation == [False]
    assert response.json()["model"] == settings.generation_model


def test_post_endpoints_require_api_key(client: TestClient):
    assert client.post("/ai/embeddings", json={"text": "claim"}).status_code == 401


def test_generate_accepts_system_and_returns_provider_metadata(
    client: TestClient, monkeypatch
):
    async def generate(system, prompt, settings, response_format, **options):
        assert system == "Judge claim quality"
        assert prompt == "Review this"
        assert response_format is None
        assert options == {"model_index": 0, "attempt": 1, "budget_ms": 300000, "validation_feedback": None}
        assert settings.ollama_model == SETTINGS.ollama_model
        return GenerateResponse(
            provider="ollama",
            model=settings.ollama_model,
            response="done",
            done=True,
        )

    monkeypatch.setattr(main, "generate_text", generate)
    response = client.post(
        "/ai/generate",
        headers=HEADERS,
        json={"system": "Judge claim quality", "prompt": "Review this"},
    )

    assert response.status_code == 200
    assert response.json() == {
        "provider": "ollama",
        "model": SETTINGS.ollama_model,
        "response": "done",
        "done": True,
        "model_index": 0,
        "attempt": 1,
        "next_model_index": None,
    }


def test_generate_keeps_prompt_only_request_compatible(client: TestClient, monkeypatch):
    async def generate(system, prompt, settings, response_format, **_options):
        assert system == ""
        assert response_format is None
        return GenerateResponse(
            provider="ollama",
            model=settings.ollama_model,
            response=prompt,
            done=True,
        )

    monkeypatch.setattr(main, "generate_text", generate)

    response = client.post(
        "/ai/generate",
        headers=HEADERS,
        json={"prompt": "Review this"},
    )

    assert response.status_code == 200
    assert response.json()["response"] == "Review this"


def test_generate_accepts_whole_paper_prompt(client: TestClient, monkeypatch):
    async def generate(_system, prompt, settings, _response_format, **_options):
        return GenerateResponse(
            provider="ollama",
            model=settings.ollama_model,
            response=str(len(prompt)),
            done=True,
        )

    monkeypatch.setattr(main, "generate_text", generate)

    accepted = client.post(
        "/ai/generate",
        headers=HEADERS,
        json={"prompt": "x" * 48000},
    )
    rejected = client.post(
        "/ai/generate",
        headers=HEADERS,
        json={"prompt": "x" * 48001},
    )

    assert accepted.status_code == 200
    assert accepted.json()["response"] == "48000"
    assert rejected.status_code == 422


def test_generate_passes_json_schema_response_format(client: TestClient, monkeypatch):
    expected = {
        "type": "json_schema",
        "json_schema": {
            "name": "section_standard",
            "strict": True,
            "schema": {
                "type": "object",
                "properties": {"passed": {"type": "boolean"}},
                "required": ["passed"],
                "additionalProperties": False,
            },
        },
    }

    async def generate(_system, _prompt, settings, response_format, **_options):
        assert response_format == expected
        return GenerateResponse(
            provider="remote",
            model=settings.generation_model or "test-model",
            response='{"passed":true}',
            done=True,
        )

    monkeypatch.setattr(main, "generate_text", generate)

    response = client.post(
        "/ai/generate",
        headers=HEADERS,
        json={"prompt": "Review this", "response_format": expected},
    )

    assert response.status_code == 200


@pytest.mark.parametrize(
    ("failure", "expected_status"),
    [
        (GenerationConfigurationError("bad config"), 503),
        (GenerationUnavailableError("offline"), 503),
        (GenerationRateLimitError("rate limited"), 429),
        (GenerationInvalidResponseError("malformed"), 502),
    ],
)
def test_generation_errors_map_to_gateway_status(failure, expected_status, monkeypatch):
    async def generate(*_, **_options):
        raise failure

    monkeypatch.setattr(main, "generate_text", generate)
    with TestClient(main.app, raise_server_exceptions=False) as client:
        response = client.post(
            "/ai/generate",
            headers=HEADERS,
            json={"prompt": "Review this"},
        )

    assert response.status_code == expected_status
    assert response.json() == {"detail": str(failure), "code": failure.code}


def test_single_and_batch_embeddings_preserve_order(client: TestClient, monkeypatch):
    async def embed(texts, _):
        return [[float(index)] for index, _text in enumerate(texts)]

    monkeypatch.setattr(main, "generate_embeddings", embed)

    single = client.post("/ai/embeddings", headers=HEADERS, json={"text": "one"})
    batch = client.post(
        "/ai/embeddings/batch",
        headers=HEADERS,
        json={"texts": ["one", "two"]},
    )

    assert single.json() == {"embedding": [0.0]}
    assert batch.json() == {"embeddings": [[0.0], [1.0]]}


def test_batch_rejects_more_than_64_texts(client: TestClient):
    response = client.post(
        "/ai/embeddings/batch",
        headers=HEADERS,
        json={"texts": ["chunk"] * 65},
    )
    assert response.status_code == 422


def test_extract_returns_bundle(monkeypatch):
    async def create_bundle(payload, _, destination):
        assert payload.filename == "paper.pdf"
        with zipfile.ZipFile(destination, "w") as archive:
            archive.writestr(
                "extraction.json",
                json.dumps({
                    "blocks": [
                        {"type": "heading", "text": "Paper", "level": 1, "caption": None}
                    ],
                    "images": ["images/figure.jpg"],
                }),
            )
            archive.writestr("document.md", "# Paper\n\n![](images/figure.jpg)")
            archive.writestr("images/figure.jpg", b"jpeg")

    monkeypatch.setattr(main, "create_extraction_bundle", create_bundle)
    with TestClient(main.app, raise_server_exceptions=False) as client:
        response = client.post(
            "/extract",
            headers=HEADERS,
            json={
                "filename": "paper.pdf",
                "download_url": "https://storage.test/paper.pdf?signature=test",
            },
        )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/zip")
    with zipfile.ZipFile(io.BytesIO(response.content)) as archive:
        assert archive.namelist() == [
            "extraction.json",
            "document.md",
            "images/figure.jpg",
        ]
        assert archive.read("document.md") == b"# Paper\n\n![](images/figure.jpg)"
        assert archive.read("images/figure.jpg") == b"jpeg"


def test_extract_rejects_removed_request_fields(client: TestClient):
    response = client.post(
        "/extract",
        headers=HEADERS,
        json={
            "document_id": str(uuid4()),
            "filename": "paper.pdf",
            "content_type": "application/pdf",
            "download_url": "https://storage.test/paper.pdf",
        },
    )

    assert response.status_code == 422


def test_extract_rejects_untrusted_download_host(client: TestClient):
    response = client.post(
        "/extract",
        headers=HEADERS,
        json={
            "filename": "paper.pdf",
            "download_url": "https://evil.test/paper.pdf",
        },
    )
    assert response.status_code == 422


def test_download_failure_logs_host_and_status_without_secret_url(
    monkeypatch, tmp_path, caplog
):
    real_client = httpx.AsyncClient
    transport = httpx.MockTransport(
        lambda request: httpx.Response(404, request=request)
    )
    monkeypatch.setattr(
        extraction.httpx,
        "AsyncClient",
        lambda **kwargs: real_client(transport=transport, **kwargs),
    )
    caplog.set_level(logging.WARNING, logger=extraction.__name__)

    with pytest.raises(ExtractionUnavailableError):
        asyncio.run(extraction._download(
            "https://storage.test/paper.pdf?token=secret-token",
            tmp_path / "paper.pdf",
            1024,
        ))

    assert (
        "source_download_failed host=storage.test status=404 "
        "error_type=HTTPStatusError"
    ) in caplog.text
    assert "secret-token" not in caplog.text


def test_extraction_routes_pdf_to_configured_mineru(monkeypatch):
    received = {}
    blocks = (SimpleNamespace(type="paragraph", text="Extracted", level=None, caption=None),)

    async def download(*_):
        return None

    async def mineru(*args):
        received["args"] = args
        return ExtractionWorkProduct(
            ExtractedDocument("# Extracted", blocks),
        )

    monkeypatch.setattr("app.extraction._download", download)
    monkeypatch.setattr("app.extraction.extract_with_mineru", mineru)
    payload = ExtractRequest(
        filename="paper.pdf",
        download_url="https://storage.test/paper.pdf",
    )

    settings = SETTINGS.model_copy(update={
        "mineru_command": "C:/tools/mineru.exe",
        "mineru_backend": "pipeline",
    })
    result = asyncio.run(extract_from_url(payload, settings))
    assert result.markdown == "# Extracted"
    assert result.blocks == blocks
    assert received["args"][-2:] == ("C:/tools/mineru.exe", "pipeline")


def test_mineru_logs_output_before_stream_ends(caplog):
    async def stream() -> bytes:
        reader = asyncio.StreamReader()
        task = asyncio.create_task(extraction._stream_mineru_output(reader))
        reader.feed_data(b"Processing page 1\n")
        await asyncio.sleep(0)
        assert "MinerU: Processing page 1" in caplog.text
        reader.feed_eof()
        return await task

    with caplog.at_level(logging.INFO, logger=extraction.__name__):
        assert asyncio.run(stream()) == b"Processing page 1\n"


def test_extraction_routes_markdown_to_structured_blocks(monkeypatch):
    async def download(_, destination, __):
        destination.write_text(
            "# Methods\n\n"
            "A result paragraph.\n\n"
            "| Metric | Value |\n"
            "| --- | --- |\n"
            "| Recall | 0.91 |\n\n"
            "# References\n\n"
            "Smith 2024",
            encoding="utf-8",
        )

    monkeypatch.setattr("app.extraction._download", download)
    payload = ExtractRequest(
        filename="source.markdown",
        download_url="https://storage.test/source.markdown",
    )

    result = asyncio.run(extract_from_url(payload, SETTINGS))

    assert result.markdown.startswith("# Methods")
    assert [block.type for block in result.blocks] == [
        "heading",
        "paragraph",
        "table",
        "reference",
        "reference",
    ]
    assert result.blocks[2].text.endswith("| Recall | 0.91 |")


def test_extraction_rejects_non_utf8_markdown(monkeypatch):
    async def download(_, destination, __):
        destination.write_bytes(b"\xff\xfe\x00")

    monkeypatch.setattr("app.extraction._download", download)
    payload = ExtractRequest(
        filename="source.md",
        download_url="https://storage.test/source.md",
    )

    with pytest.raises(ExtractionError, match="UTF-8"):
        asyncio.run(extract_from_url(payload, SETTINGS))


def test_markdown_parser_preserves_setext_lists_and_fenced_code():
    result = extraction._document_from_markdown(
        "Results\n"
        "=======\n\n"
        "- first\n"
        "- second\n\n"
        "```python\n"
        "print('ok')\n"
        "```"
    )

    assert [block.type for block in result.blocks] == ["heading", "list", "code"]
    assert result.blocks[0].level == 1
    assert result.blocks[1].text == "- first\n- second"
    assert result.blocks[2].text == "print('ok')"


def test_extraction_routes_docx_to_structured_blocks(monkeypatch):
    async def download(_, destination, __):
        document = WordDocument()
        document.add_heading("Methods", level=1)
        document.add_paragraph("A result paragraph.")
        document.add_paragraph("First item", style="List Bullet")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Recall"
        table.rows[1].cells[1].text = "0.91"
        document.add_heading("References", level=1)
        document.add_paragraph("Smith 2024")
        document.save(destination)

    monkeypatch.setattr("app.extraction._download", download)
    payload = ExtractRequest(
        filename="source.docx",
        download_url="https://storage.test/source.docx",
    )

    result = asyncio.run(extract_from_url(payload, SETTINGS))

    assert result.markdown.startswith("# Methods")
    assert [block.type for block in result.blocks] == [
        "heading",
        "paragraph",
        "list",
        "table",
        "reference",
        "reference",
    ]
    assert "| Recall | 0.91 |" in result.blocks[3].text


def test_extraction_rejects_invalid_docx(monkeypatch):
    async def download(_, destination, __):
        destination.write_bytes(b"not-a-docx")

    monkeypatch.setattr("app.extraction._download", download)
    payload = ExtractRequest(
        filename="source.docx",
        download_url="https://storage.test/source.docx",
    )

    with pytest.raises(ExtractionError, match="DOCX file is invalid"):
        asyncio.run(extract_from_url(payload, SETTINGS))


def test_extraction_rejects_docx_expanding_beyond_limit(monkeypatch):
    async def download(_, destination, __):
        with ZipFile(destination, "w") as archive:
            archive.writestr("[Content_Types].xml", b"x" * 101)

    monkeypatch.setattr("app.extraction._download", download)
    payload = ExtractRequest(
        filename="source.docx",
        download_url="https://storage.test/source.docx",
    )
    settings = SETTINGS.model_copy(update={"max_download_bytes": 100})

    with pytest.raises(ExtractionError, match="DOCX content exceeds"):
        asyncio.run(extract_from_url(payload, settings))


def test_extract_rejects_tex_with_422(client: TestClient):
    response = client.post(
        "/extract",
        headers=HEADERS,
        json={
            "filename": "source.tex",
            "download_url": "https://storage.test/source.tex",
        },
    )

    assert response.status_code == 422
    assert response.json() == {
        "detail": "only PDF, DOCX, and Markdown files are supported"
    }


def test_mineru_reads_markdown_and_normalizes_blocks(tmp_path):
    output_dir = tmp_path / "input" / "hybrid_auto"
    markdown_path = output_dir / "input.md"
    markdown_path.parent.mkdir(parents=True)
    markdown_path.write_text("# Extracted", encoding="utf-8")
    (output_dir / "input_content_list.json").write_text(
        json.dumps([
            {"type": "text", "text": "Results", "text_level": 2},
            {"type": "text", "text": "A result paragraph."},
            {
                "type": "table",
                "table_caption": ["Table 1"],
                "table_body": (
                    "<table><thead><tr><th>A</th><th>B</th></tr></thead>"
                    "<tbody><tr><td>1</td><td>2</td></tr></tbody></table>"
                ),
            },
            {"type": "image", "image_caption": ["Figure 1. Architecture"]},
            {"type": "text", "text": "References", "text_level": 1},
            {"type": "list", "sub_type": "ref_text", "text": "Smith 2024"},
        ]),
        encoding="utf-8",
    )

    read_output = getattr(extraction, "_read_mineru_output", None)
    assert read_output is not None
    result = read_output(tmp_path, "input")

    assert result.document.markdown == "# Extracted"
    assert [block.type for block in result.document.blocks] == [
        "heading",
        "paragraph",
        "table",
        "figure_caption",
        "reference",
        "reference",
    ]
    assert result.document.blocks[2].caption == "Table 1"
    assert result.document.blocks[2].text == "| A | B |\n| --- | --- |\n| 1 | 2 |"


def test_mineru_flat_headings_use_validated_local_hierarchy(monkeypatch):
    blocks = tuple(
        extraction.ExtractionBlock("heading", text, level)
        for text, level in [
            ("Paper title", 1),
            ("Introduction", 2),
            ("Study design", 2),
            ("Results", 2),
            ("Limitations", 2),
        ]
    )

    async def hierarchy(_, prompt, __, **options):
        assert "exactly 5 headings" in prompt
        assert '"index": 4' in prompt
        assert '"fixed_level": 1' in prompt
        assert "current_level" not in prompt
        assert options["budget_ms"] == 30000
        options["validate"](json.dumps({"levels": [1, 2, 3, 2, 3]}))
        return GenerateResponse(
            provider="ollama",
            model="test-model",
            response=json.dumps({"levels": [1, 2, 3, 2, 3]}),
            done=True,
        )

    monkeypatch.setattr(extraction, "generate_text", hierarchy)
    product = ExtractionWorkProduct(ExtractedDocument("# Paper title", blocks))

    enriched = asyncio.run(extraction._enrich_mineru_hierarchy(product, SETTINGS))

    assert [block.level for block in enriched.document.blocks] == [1, 2, 3, 2, 3]


def test_mineru_keeps_flat_levels_when_hierarchy_is_invalid(monkeypatch):
    blocks = (
        extraction.ExtractionBlock("heading", "Paper title", 1),
        extraction.ExtractionBlock("heading", "Introduction", 2),
        extraction.ExtractionBlock("heading", "Study design", 2),
    )

    async def hierarchy(*_, **_options):
        return GenerateResponse(
            provider="ollama",
            model="test-model",
            response=json.dumps({"levels": [1, 4, 2]}),
            done=True,
        )

    monkeypatch.setattr(extraction, "generate_text", hierarchy)
    product = ExtractionWorkProduct(ExtractedDocument("# Paper title", blocks))

    enriched = asyncio.run(extraction._enrich_mineru_hierarchy(product, SETTINGS))

    assert enriched is product


def test_mineru_bundle_includes_referenced_image(tmp_path):
    output_dir = tmp_path / "output"
    document_dir = output_dir / "paper"
    image_path = document_dir / "images" / "figure.jpg"
    image_path.parent.mkdir(parents=True)
    image_path.write_bytes(b"jpeg")
    (document_dir / "paper.md").write_text(
        "# Result\n\n![](images/figure.jpg)",
        encoding="utf-8",
    )
    (document_dir / "paper_content_list.json").write_text(
        json.dumps([
            {"type": "text", "text": "Result", "text_level": 1},
            {"type": "image", "img_path": "images/figure.jpg"},
        ]),
        encoding="utf-8",
    )

    product = extraction._read_mineru_output(output_dir, "paper")

    assert product.document.images == ("images/figure.jpg",)
    assert product.image_files == (("images/figure.jpg", image_path),)


def test_mineru_rejects_image_path_escape(tmp_path):
    output_dir = tmp_path / "output"
    document_dir = output_dir / "paper"
    document_dir.mkdir(parents=True)
    (document_dir / "paper.md").write_text("# Result", encoding="utf-8")
    (document_dir / "paper_content_list.json").write_text(
        json.dumps([
            {"type": "text", "text": "Result", "text_level": 1},
            {"type": "image", "img_path": "../secret.jpg"},
        ]),
        encoding="utf-8",
    )

    with pytest.raises(ExtractionUnavailableError, match="image path"):
        extraction._read_mineru_output(output_dir, "paper")


def test_mineru_requires_content_list(tmp_path):
    output_dir = tmp_path / "input" / "hybrid_auto"
    output_dir.mkdir(parents=True)
    (output_dir / "input.md").write_text("# Extracted", encoding="utf-8")

    read_output = getattr(extraction, "_read_mineru_output", None)
    assert read_output is not None
    with pytest.raises(ExtractionUnavailableError, match="content list"):
        read_output(tmp_path, "input")


def test_extraction_block_level_is_only_valid_for_headings():
    with pytest.raises(ValueError):
        ExtractionBlock(type="heading", text="Methods")
    with pytest.raises(ValueError):
        ExtractionBlock(type="paragraph", text="Body", level=2)


def test_settings_reads_mineru_command(monkeypatch):
    monkeypatch.setenv("MINERU_COMMAND", "C:/tools/mineru.exe")
    monkeypatch.setenv("MINERU_BACKEND", "pipeline")

    settings = Settings.from_env()

    assert getattr(settings, "mineru_command", None) == "C:/tools/mineru.exe"
    assert settings.mineru_backend == "pipeline"


def test_settings_reads_generation_provider_configuration(monkeypatch):
    monkeypatch.setenv("GENERATION_PROVIDER", "remote")
    monkeypatch.setenv("GENERATION_API_KEY", "router-secret")
    monkeypatch.setenv(
        "GENERATION_BASE_URL",
        "https://gateway.test/v1/",
    )
    monkeypatch.setenv("GENERATION_MODEL", "nemotron-test")
    monkeypatch.setenv("GENERATION_FALLBACK_MODELS", '["gemma-test", "super-test"]')
    monkeypatch.setenv(
        "GENERATION_EXTRA_BODY",
        '{"reasoning":{"effort":"none"}}',
    )
    monkeypatch.setenv("OLLAMA_MODEL", "qwen3.5:9b")

    settings = Settings.from_env()

    assert settings.generation_provider == "remote"
    assert settings.generation_api_key == "router-secret"
    assert settings.generation_base_url == "https://gateway.test/v1"
    assert settings.generation_model == "nemotron-test"
    assert settings.generation_fallback_models == ["gemma-test", "super-test"]
    assert settings.generation_extra_body == {
        "reasoning": {"effort": "none"}
    }
    assert settings.ollama_model == "qwen3.5:9b"


def test_generate_continuation_contract(client, monkeypatch):
    async def generate(*_, **options):
        assert options == {"model_index": 1, "attempt": 2, "budget_ms": 90000,
                           "validation_feedback": "Quote not found"}
        return GenerateResponse(provider="remote", model="actual", response="{}", done=True,
                                model_index=2, attempt=1)

    monkeypatch.setattr(main, "generate_text", generate)
    response = client.post("/ai/generate", headers=HEADERS, json={
        "prompt": "Review", "model_index": 1, "attempt": 2, "budget_ms": 90000,
        "validation_feedback": "Quote not found",
    })
    assert response.status_code == 200
    assert (response.json()["model_index"], response.json()["attempt"]) == (2, 1)


def test_invalid_schema_returns_422_without_calling_provider(client, monkeypatch):
    async def unexpected(*_, **__):
        pytest.fail("Invalid schema must not call the provider")

    monkeypatch.setattr(OllamaGenerationProvider, "generate", unexpected)
    response = client.post("/ai/generate", headers=HEADERS, json={
        "prompt": "Review", "response_format": {"type": "json_schema", "json_schema": {
            "name": "result", "schema": {"$ref": "https://private.test/schema"},
        }},
    })
    assert response.status_code == 422
    assert response.json()["code"] == "INVALID_GENERATION_REQUEST"
    assert "private.test" not in response.text


def test_rate_limit_preserves_safe_retry_after(client, monkeypatch):
    async def generate(*_, **__):
        raise GenerationRateLimitError("Provider rate limit exceeded", retry_after=12.5)

    monkeypatch.setattr(main, "generate_text", generate)
    response = client.post("/ai/generate", headers=HEADERS, json={"prompt": "Review"})
    assert response.status_code == 429 and response.headers["Retry-After"] == "13"


def test_hierarchy_releases_local_slot_before_remote_call(monkeypatch):
    from app.ollama_client import generate_embeddings

    product = ExtractionWorkProduct(ExtractedDocument("# Title", (
        extraction.ExtractionBlock("heading", "Title", 1),
        extraction.ExtractionBlock("heading", "Introduction", 2),
        extraction.ExtractionBlock("heading", "Methods", 2),
    )))

    async def download(*_):
        pass

    async def mineru(*_):
        return product

    async def hierarchy(*_, **options):
        # This would deadlock with max_concurrent=1 if extraction still held its slot.
        vectors = await asyncio.wait_for(generate_embeddings(["synthetic"], SETTINGS), 0.5)
        assert vectors == [[1.0]]
        raise GenerationUnavailableError("Generation batch deadline exceeded")

    real_client = httpx.AsyncClient
    transport = httpx.MockTransport(lambda _: httpx.Response(200, json={"embeddings": [[1.0]]}))
    monkeypatch.setattr(httpx, "AsyncClient", lambda **kwargs: real_client(transport=transport, **kwargs))
    monkeypatch.setattr(limits, "local_gate", limits.ModelCallGate(1))
    monkeypatch.setattr(limits, "generation_gate", limits.ModelCallGate(1))
    monkeypatch.setattr(extraction, "_download", download)
    monkeypatch.setattr(extraction, "extract_with_mineru", mineru)
    monkeypatch.setattr(extraction, "generate_text", hierarchy)

    async def exercise():
        async with limits.generation_gate.slot():
            return await extract_from_url(ExtractRequest(
                filename="paper.pdf", download_url="https://storage.test/paper.pdf",
            ), SETTINGS)

    assert asyncio.run(exercise()) is product.document


def test_mineru_cancellation_stops_process_and_drains_streams(monkeypatch, tmp_path):
    class Process:
        returncode = None
        killed = False

        def __init__(self):
            self.stdout = asyncio.StreamReader()
            self.stderr = asyncio.StreamReader()
            self.started = asyncio.Event()
            self.finished = asyncio.Event()

        async def wait(self):
            self.started.set()
            await self.finished.wait()

        def kill(self):
            self.killed = True
            self.returncode = -1
            self.stdout.feed_eof()
            self.stderr.feed_eof()
            self.finished.set()

    async def exercise():
        process = Process()

        async def start(*_, **__):
            return process

        monkeypatch.setattr(asyncio, "create_subprocess_exec", start)
        task = asyncio.create_task(extraction.extract_with_mineru(
            tmp_path / "paper.pdf", tmp_path / "out", 60, "mineru", "pipeline",
        ))
        await process.started.wait()
        task.cancel()
        with pytest.raises(asyncio.CancelledError):
            await task
        assert process.killed and process.stdout.at_eof() and process.stderr.at_eof()

    asyncio.run(exercise())
